import re
import io
import os
import logging
from docx import Document
from docx.shared import Inches
from bs4 import BeautifulSoup

logger = logging.getLogger('DOCX_BINDER')

def build_paragraph(paragraph, elem):
    for tag in elem.children:
        if not tag.name:
            paragraph.add_run(str(tag))
        elif re.match(r'b|strong|label|kdb', tag.name):
            paragraph.add_run(tag.text).bold = True
        elif re.match(r'i|em|pre|code|cite', tag.name):
            paragraph.add_run(tag.text).italic = True
        else:
            paragraph.add_run(tag.text)
        # end if
        paragraph.add_run(' ')
    # end for
# end def

def make_chapter(book, chapter):
    soup = BeautifulSoup(chapter['body'], 'lxml')
    for tag in soup.find('body').children:
        if not tag.name:
            book.add_paragraph(str(tag))
        elif re.match(r'h\d', tag.name):
            book.add_heading(tag.text, int(tag.name[1]))
        elif tag.name == 'p':
            build_paragraph(book.add_paragraph(), tag)
        # end if
    # end for
    book.add_page_break()
# end def


def bind_docx_book(app, chapters, volume=''):
    book_title = (app.crawler.novel_title + ' ' + volume).strip()
    logger.debug('Binding %s.docx', book_title)

    # Create book
    book = Document()
    book.core_properties.language = 'en'
    book.core_properties.title = book_title
    book.core_properties.category = 'Lightnovel'
    book.core_properties.author = 'Lightnovel Crawler'
    book.core_properties.identifier = app.output_path + volume

    # Create intro page
    book.add_heading(app.crawler.novel_title or 'N/A', 0)
    book.add_paragraph((app.crawler.novel_author or '').replace(':', ': '))
    book.add_paragraph('\r\n' * 5)
    p = book.add_paragraph('Generated by ')
    p.add_run('Lightnovel Crawler').bold = True
    p.add_run('\r\nhttps://github.com/dipu-bd/lightnovel-crawler')
    # book.add_picture(app.book_cover)  # , width=Inches(1.25))
    book.add_page_break()

    # Create chapters
    for chapter in chapters:
        make_chapter(book, chapter)
    # end for

    # Save epub file
    epub_path = os.path.join(app.output_path, 'docx')
    file_path = os.path.join(epub_path, book_title + '.docx')
    logger.debug('Writing %s', file_path)
    os.makedirs(epub_path, exist_ok=True)
    book.save(file_path)
    logger.warn('Created: %s.docx', book_title)
    return file_path
# end def


def make_docx(app, data):
    docx_files = []
    for vol in data:
        if len(data[vol]) > 0:
            book = bind_docx_book(
                app,
                volume=vol,
                chapters=data[vol],
            )
            docx_files.append(book)
        # end if
    # end for
    return docx_files
# end def
